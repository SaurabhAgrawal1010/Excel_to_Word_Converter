import pandas as pd
from docx import Document

document = Document()
file_name = 'Path-to-directory/test.xlsx'

exc = pd.ExcelFile(file_name)
        
for s in exc.sheet_names:
    try:
        df = pd.read_excel(exc, 
                           sheetname=s
                           )
        #del df['Description']
        df = df.fillna('')
        r = df.shape[0]
        c = df.shape[1]
        #print(r,c)
        #print(list(df.columns.values))
        
        #print(df)
        t = document.add_table(rows=(r+1), cols=c, style='Table Grid')
        #t.style = 'TableGrid'
        
        # add the header rows.
        for j in range(df.shape[-1]):
            header_text = df.columns[j]
            row = t.rows[0]
            header_text_formatted = row.cells[j].paragraphs[0].add_run(header_text)
            header_text_formatted.bold = True
            #header_text_formatted.font.name = 'Arial'
            #header_text_formatted.font.size = Pt(8)

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                body_text = str(df.values[i,j])
                row = t.rows[i+1]
                body_text_formatted = row.cells[j].paragraphs[0].add_run(body_text)
                #body_text_formatted.bold = True
                #body_text_formatted.font.name = 'Arial'
                #body_text_formatted.font.size = Pt(8)

        document.save('Path-to-directory/test.docx')
            
    except Exception as e:
        print(e)
